from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Codex\HarveyClassDemo")
OUT_DIR = ROOT / "deliverables"
OUT_DIR.mkdir(parents=True, exist_ok=True)

DETAILED_PATH = OUT_DIR / "LEGO_Entertainment_Walkthrough_Detailed.docx"
CUE_PATH = OUT_DIR / "LEGO_Entertainment_Presenter_Cue_Sheet.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "111318"
MUTED = "5E6673"
LIGHT = "F4F6F9"
LEGO_RED = "E2231A"
LEGO_YELLOW = "FFCF00"
LEGO_BLUE = "087BC1"
LEGO_GREEN = "00A651"
PALE_YELLOW = "FFF7CE"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"


SLIDES = [
    {
        "num": 1,
        "title": "Meet the Crew",
        "time": "0:35",
        "clicks": "2 clicks",
        "purpose": "先讓觀眾掃 QR，再介紹團隊與唯一的核心問題。",
        "script": (
            "Good morning. Before we begin, please take a moment to scan the QR code on screen. "
            "[PAUSE, THEN CLICK.] Thank you. We are Harvey Yang, Olivia, Tinya, June, and Anglea. "
            "Today, we are not simply listing LEGO films, television series, and games. We want to "
            "explain how entertainment changed what LEGO could become. Our central question is whether "
            "LEGO abandoned the brick to become an entertainment company, or used entertainment to make "
            "the brick system more valuable."
        ),
        "delivery": "QR 停留 5–8 秒。第一次按鍵後看向人物，名字只說一次；不要重新介紹课程背景。",
        "transition": "To answer that question, we first need to understand what entertainment actually does for LEGO.",
        "remote": "CLICK 1 after the QR pause. CLICK 2 after the transition.",
        "open": "Good morning. Before we begin, please scan the QR code on screen.",
        "anchors": "QR → team → one strategic question",
        "number": "Five presenters; one core story",
        "close": "Did LEGO abandon the brick, or make the brick more valuable?",
    },
    {
        "num": 2,
        "title": "LEGO Story Engines",
        "time": "0:35",
        "clicks": "1 click",
        "purpose": "解释电影、电视、游戏与平台为何不是重复投资，而是承担不同任务。",
        "script": (
            "Entertainment turns LEGO bricks into a repeatable play engine. Film creates cultural reach. "
            "Television keeps characters and worlds alive. Games give players agency because they can "
            "change and rebuild the world. Digital platforms create recurring participation through "
            "communities and live updates. These media perform different jobs, but they create strategic "
            "value only when attention, characters, and new ideas eventually return to physical LEGO play."
        ),
        "delivery": "顺着 Film、TV、Games、Platforms 指过去；每项只用一个动词，不逐字读页面。",
        "transition": "We can see this mechanism more clearly through three important LEGO worlds.",
        "remote": "CLICK once after the transition.",
        "open": "Entertainment turns LEGO bricks into a repeatable play engine.",
        "anchors": "Film = reach; TV = continuity; games = agency; platforms = recurrence",
        "number": "Four media, four jobs",
        "close": "Every loop must return value to physical play.",
    },
    {
        "num": 3,
        "title": "LEGO Worlds",
        "time": "1:10",
        "clicks": "4 clicks",
        "purpose": "用三个案例预告整场演化：连续性、品牌意义、持续参与。",
        "script": (
            "These three worlds represent three stages in LEGO's entertainment evolution: continuity through "
            "television, brand meaning through film, and persistent participation through a digital platform. "
            "[CLICK: NINJAGO.] NINJAGO turned LEGO's own product characters into a renewable story world. Every "
            "season creates heroes, villains, vehicles, locations, and conflicts that can become new physical sets. "
            "[CLICK: THE LEGO MOVIE.] The LEGO Movie made LEGO's philosophy the story. Building, rebuilding, "
            "creativity, and challenging instructions became ideas a global audience could understand. "
            "[CLICK: LEGO FORTNITE.] LEGO Fortnite creates an always-on relationship. Players do not simply watch "
            "a story; they build, explore, create, and return. Together, the cases show a progression from attention, "
            "to meaning, to continuous participation. But the capability behind them began much earlier."
        ),
        "delivery": "每次弹窗变化后停半秒再说。眼睛看观众，电脑画面只是证明，不要讲 achievement 三点。",
        "transition": "But the capability behind these worlds began much earlier.",
        "remote": "CLICK to open NINJAGO; CLICK for Movie; CLICK for Fortnite; CLICK after the final sentence to continue.",
        "open": "Three worlds show three stages of LEGO's entertainment evolution.",
        "anchors": "NINJAGO = continuity; Movie = meaning; Fortnite = participation",
        "number": "Four remote presses on this page",
        "close": "The capability behind them began much earlier.",
    },
    {
        "num": 4,
        "title": "Original Advantage",
        "time": "0:30",
        "clicks": "1 click",
        "purpose": "把娱乐定位为既有 System in Play 的延伸，而不是完全无关的新业务。",
        "script": (
            "LEGO's original advantage was not one hit toy. It was a system. The compatible brick system arrived "
            "in 1958. Minifigures added identity in 1978. Star Wars demonstrated licensed translation in 1999. "
            "NINJAGO created an owned world in 2011, and LEGO Fortnite added a persistent digital world in 2023. "
            "The brick created compatibility, the minifigure created identity, and entertainment made both repeatable across media."
        ),
        "delivery": "沿时间线从左到右移动手势，1958 与 1978 稍停，后面年份加快。",
        "transition": "The most important bridge between compatible bricks and entertainment was the minifigure.",
        "remote": "CLICK once after the transition.",
        "open": "LEGO's original advantage was not one hit toy. It was a system.",
        "anchors": "1958 compatibility → 1978 identity → worlds",
        "number": "1958 / 1978 / 1999 / 2011 / 2023",
        "close": "The minifigure was the bridge to entertainment.",
    },
    {
        "num": 5,
        "title": "Minifigures Become Stories",
        "time": "0:40",
        "clicks": "1 click",
        "purpose": "把抽象能力落实到一个人、一个设计决定与一个结果。",
        "script": (
            "The bridge had a designer: Jens N. Knudsen, who worked at LEGO from 1968 to 2000. His decision was "
            "not simply to make a smaller figure. It was to create one articulated character neutral enough to play "
            "any role. The 1974 figure was brick-built. The 1975 figure was a fixed stage extra. In 1978, the modern "
            "minifigure could move, travel between sets, and become a hero, parent, astronaut, or villain. Buildings "
            "gained inhabitants, characters gained identity, and play gained continuing stories. LEGO did not enter "
            "entertainment in 1978; it invented the character that made entertainment possible."
        ),
        "delivery": "先指 Jens 卡片，再依次指 1974、1975、1978。最后一句放慢并看观众。",
        "transition": "LEGO first scaled this character capability by entering a story audiences already knew.",
        "remote": "CLICK once after the transition.",
        "source": "LEGO Group History, Role Play and Minifigure Timeline; Jens N. Knudsen tenure, 1968–2000.",
        "open": "The bridge had a designer: Jens N. Knudsen.",
        "anchors": "one neutral character → any role → continuing story",
        "number": "32 years at LEGO; breakthrough in 1978",
        "close": "LEGO invented the character that made entertainment possible.",
    },
    {
        "num": 6,
        "title": "Licensed Worlds",
        "time": "0:45",
        "clicks": "1 click",
        "purpose": "说明 Star Wars 的关键价值是能力学习，而不只是授权销售。",
        "script": (
            "That next decision came from Peter Eio, President of LEGO Americas. In late 1997, he proposed a "
            "partnership with Lucasfilm even though LEGO had never licensed an outside story. Senior leaders debated "
            "the brand fit for about six months: could a company built around open-ended play work with a story about war? "
            "LEGO signed its first major IP agreement in 1998 and launched LEGO Star Wars in 1999. Star Wars supplied "
            "recognition, but LEGO supplied the translation: ships became buildable systems and fans could break, remix, "
            "and replay the story. The licence became a training ground for converting fandom into participation."
        ),
        "delivery": "左边产品图只作视觉证据；重点指 Peter Eio 卡片与四步时间线。",
        "transition": "But expansion did not automatically create value. In 2004, complexity caught up with LEGO.",
        "remote": "CLICK once after the transition.",
        "source": "The LEGO Group, LEGO Star Wars 25th Anniversary (2024); Peter Eio's account of the 1997 proposal and internal debate.",
        "open": "The bet that broke LEGO's purist rule came from Peter Eio.",
        "anchors": "outside IP → internal debate → LEGO translation",
        "number": "Late 1997; six-month debate; 1999 launch",
        "close": "Licensing became LEGO's training ground.",
    },
    {
        "num": 7,
        "title": "The 2004 Crisis",
        "time": "0:45",
        "clicks": "1 click",
        "purpose": "避免错误归因：娱乐不是危机的唯一原因，失控复杂度才是警告。",
        "script": (
            "The 2004 crisis is the warning inside this growth story. LEGO reported revenue of about 6.7 billion "
            "Danish kroner, but a net loss of roughly 1.9 billion. After 25 years as CEO, Kjeld Kirk Kristiansen handed "
            "operational leadership to Jørgen Vig Knudstorp, who had joined LEGO in 2001 and became CEO at 35. The plan "
            "was to cut product complexity, release capital from heavy assets, and use partners for specialist execution. "
            "By 2005, LEGO reported 702 million kroner in profit before tax, although restructuring and asset transactions "
            "also affected that result. Entertainment was not the problem. Expansion without a clear return to the core was."
        ),
        "delivery": "先看左侧亏损，再用手势完成 Kjeld → Jørgen 的领导交接。不要把 2005 结果说成单一措施造成。",
        "transition": "The recovery question was whether LEGO could own worlds without owning every production capability.",
        "remote": "CLICK once after the transition.",
        "source": "The LEGO Group Annual Reports 2004–2005; causality stated cautiously.",
        "open": "The 2004 crisis is the warning inside this growth story.",
        "anchors": "complexity → leadership handoff → refocus and partner",
        "number": "−DKK 1.9B; 2005 +DKK 702M pre-tax",
        "close": "Expansion without a return to the core was the problem.",
    },
    {
        "num": 8,
        "title": "NINJAGO",
        "time": "0:50",
        "clicks": "1 click",
        "purpose": "说明 LEGO 如何从翻译别人的世界，升级为经营自己的长期 IP。",
        "script": (
            "NINJAGO answered that question. Tommy Andreasen and the LEGO creative team spent about 18 months studying "
            "enduring play patterns, then combined ninja, elemental powers, vehicles, and serialized conflict. Dan and "
            "Kevin Hageman helped give the product idea a continuing mythology for television. The result is more than "
            "500 sets, over 200 episodes, and 15 years of owned IP. Television provides a rhythm: each season can introduce "
            "new heroes, villains, locations, vehicles, and conflicts that are immediately buildable. Star Wars taught LEGO "
            "how to translate a world. NINJAGO proved LEGO could own the world being translated."
        ),
        "delivery": "影片自动静音播放，不需要操作。先让画面跑两秒，再说人物与 18 months；最后指三组数字。",
        "transition": "Television creates continuity. Film performs a different job: it can change what the brand means.",
        "remote": "CLICK once after the transition; do not click the video.",
        "source": "The LEGO Group, NINJAGO 15th Anniversary media release and official footage (2026).",
        "open": "NINJAGO shows what happened when LEGO moved from borrowed to owned worlds.",
        "anchors": "18-month research → TV mythology → renewable owned IP",
        "number": "500+ sets; 200+ episodes; 15 years",
        "close": "NINJAGO proved LEGO could own the world.",
    },
    {
        "num": 9,
        "title": "The LEGO Movie",
        "time": "0:45",
        "clicks": "1 click",
        "purpose": "把电影价值从票房提升到品牌意义，同时强调伙伴承担制作能力。",
        "script": (
            "The LEGO Movie did not begin as a long advertisement. Producer Dan Lin pitched a theatrical film in 2008. "
            "LEGO leaders, including Jill Wilfert, protected the brand voice and the principle of open-ended play. Phil Lord "
            "and Chris Miller then built the 2014 film around a conflict every LEGO user understands: follow the instructions "
            "or rebuild the rules. LEGO reported 15 percent consumer-sales growth in 2014 and described the movie as a "
            "significant contributor, not the only cause. By 2020, the LEGO film franchise had produced about 1.1 billion "
            "dollars in combined box office. The studio made the film; LEGO kept the meaning that anything can be rebuilt."
        ),
        "delivery": "沿右侧 2008 → LEGO → 2014 讲决定链；15% 必须加上 ‘significant contributor, not the only cause’。",
        "transition": "Film creates a cultural moment. A live digital world can create a relationship that continues every day.",
        "remote": "CLICK once after the transition.",
        "source": "The LEGO Group 2014 Annual Results; LEGO–Universal film partnership announcement (2020).",
        "open": "The LEGO Movie made the philosophy of play the plot.",
        "anchors": "2008 pitch → protect brand voice → rebuild the rules",
        "number": "+15% consumer sales; ≈$1.1B franchise box office",
        "close": "The studio made the film; LEGO kept the meaning.",
    },
    {
        "num": 10,
        "title": "LEGO Fortnite",
        "time": "0:50",
        "clicks": "1 click",
        "purpose": "说明平台带来的不是一次曝光，而是持续参与与 screen-to-shelf 双向循环。",
        "script": (
            "LEGO Fortnite changes the relationship again. In 2022, LEGO CEO Niels B. Christiansen and Epic CEO Tim Sweeney "
            "committed to a long-term partnership for safe digital play. LEGO did not try to operate the platform alone; Epic "
            "provided scale and creation tools while LEGO supplied the play system and child-safety principles. The free world "
            "launched inside Fortnite in December 2023. By the end of 2024 it had engaged more than 87 million players, and its "
            "digital characters, places, and items had returned to shelves as physical sets. The strongest platform loop works "
            "both ways: screen to shelf, and shelf back to screen."
        ),
        "delivery": "左侧图片会自动轮播，不要等待特定画面。右侧先讲两位 CEO，再讲 2023、2024 与 87M+。",
        "transition": "The cases succeed in different ways, so each medium needs a clearly defined job.",
        "remote": "CLICK once after the transition; the image loop is automatic.",
        "source": "Epic Games × LEGO partnership (2022); launch (2023); physical sets and LEGO Group 2024 Annual Results.",
        "open": "LEGO Fortnite turned a campaign into a place people could return to.",
        "anchors": "partner for platform → persistent world → screen-to-shelf",
        "number": "December 2023; 87M+ players; physical sets in 2024",
        "close": "The platform loop works in both directions.",
    },
    {
        "num": 11,
        "title": "Four Media Jobs",
        "time": "0:35",
        "clicks": "1 click",
        "purpose": "给组合投资一个统一的评价框架，避免用同一指标评价所有媒体。",
        "script": (
            "These media should not be judged by the same standard. Film provides burst reach and turns LEGO play into "
            "culture. Television provides continuity and allows characters to compound over time. Games provide agency because "
            "players can remake the world. Platforms provide recurrence by connecting content, community, insight, and products. "
            "No single medium performs all four jobs equally well. Their strategic value comes from coordination around the same IP "
            "and the same play system."
        ),
        "delivery": "按红、黄、蓝、绿四张卡从左到右扫过去；每张只说 role，不再复述案例。",
        "transition": "Once these four jobs connect, LEGO is no longer operating a simple product pipeline.",
        "remote": "CLICK once after the transition.",
        "open": "Different media should not be judged by the same standard.",
        "anchors": "reach / continuity / agency / recurrence",
        "number": "One IP system, four complementary jobs",
        "close": "Their value comes from coordination.",
    },
    {
        "num": 12,
        "title": "The Play Ecosystem",
        "time": "0:40",
        "clicks": "1 click",
        "purpose": "讲清全场最重要的因果循环：故事与参与如何回到新产品。",
        "script": (
            "This is the central transformation. LEGO once operated mainly as a product pipeline: design a set, manufacture it, "
            "sell it, and begin again. Entertainment turns that pipeline into a loop. Bricks inspire characters and stories. Stories "
            "generate media, participation, and community. Participation produces insight and new ideas. Those ideas return as new sets "
            "and new play experiences. Entertainment therefore changes LEGO from selling isolated products to renewing relationships. "
            "The loop still depends on shared core resources: brand trust, compatibility, minifigure grammar, design capability, and retail reach."
        ),
        "delivery": "沿 01–06 顺序指一圈，最后手指停在底部 SHARED CORE。这里不要加新例子。",
        "transition": "The loop is attractive, but LEGO still has to decide how much of it the company should own.",
        "remote": "CLICK once after the transition.",
        "open": "This is the central transformation: a pipeline becomes a loop.",
        "anchors": "bricks → stories → participation → insight → new sets",
        "number": "Six linked stages; one shared core",
        "close": "Entertainment renews relationships, not just products.",
    },
    {
        "num": 13,
        "title": "Three Paths Forward",
        "time": "0:45",
        "clicks": "1 click",
        "purpose": "把战略选择压缩为低控制、高复杂度与生态编排三条路径。",
        "script": (
            "LEGO has three possible paths. Option A is brick-first licensing. It requires less capital, but LEGO gives up control "
            "over owned worlds and digital relationships. Option B is full vertical integration. It provides greater control, but adds "
            "expensive and volatile capabilities that are not unique to LEGO. Option C is the ecosystem-orchestrator model. LEGO owns the "
            "architecture—its worlds, play rules, brand, and product connections—while specialist partners execute film, games, and platforms. "
            "Option C gives LEGO control where it matters without repeating the complexity trap."
        ),
        "delivery": "A 与 B 各用一句；在黄色 C 卡停留时间最长。不要把 C 说成完全外包。",
        "transition": "VRIO helps explain why the orchestrator model is the strongest strategic fit.",
        "remote": "CLICK once after the transition.",
        "open": "The choice is not whether to use entertainment, but what LEGO must control.",
        "anchors": "A low control; B high complexity; C own architecture",
        "number": "Three paths; recommend C",
        "close": "Control where it matters without repeating the complexity trap.",
    },
    {
        "num": 14,
        "title": "VRIO",
        "time": "0:40",
        "clicks": "1 click",
        "purpose": "用 VRIO 解释为什么 LEGO 应控制 play architecture，而不是拥有所有制作设施。",
        "script": (
            "VRIO asks whether a capability is valuable, rare, difficult to imitate, and supported by the organization. LEGO's brand and "
            "family trust, brick and minifigure grammar, and ability to translate a world into play and products meet all four tests. These "
            "belong on the control side. Feature-film production, AAA game engineering, and live-platform operations are valuable, but they are "
            "available from capable partners and are not uniquely LEGO. LEGO should therefore control the architecture and translation into play, "
            "while renting specialist execution."
        ),
        "delivery": "先讲左边 control，再讲右边 partner；中间 OWN ↔ PARTNER 是界线，不是二选一。",
        "transition": "VRIO tells us what LEGO should control. Related diversification tells us how far it should expand.",
        "remote": "CLICK once after the transition.",
        "source": "VRIO framework: Barney (1991).",
        "open": "LEGO's durable advantage still comes from the core, not from owning a studio.",
        "anchors": "control rare architecture; rent available execution",
        "number": "V / R / I / O",
        "close": "Control the architecture; partner for execution.",
    },
    {
        "num": 15,
        "title": "Related Diversification",
        "time": "0:40",
        "clicks": "1 click",
        "purpose": "给娱乐扩张划清边界：必须共享能力并回到积木价值。",
        "script": (
            "This loop defines the boundary of related diversification. Physical bricks provide the core resources. Bricks create characters "
            "and stories. Media and platforms extend their reach and invite participation. That participation must return as new sets, mechanics, "
            "and physical play. The Fortnite strip at the bottom makes the test visible: a digital world produced physical characters and sets. "
            "When resources are shared and stories become buildable, LEGO gains economies of scope. An isolated content bet with no shared capability "
            "and no route back to physical play is over the line."
        ),
        "delivery": "先扫上方四步，再对比 RELATED／OVER THE LINE，最后用 SCREEN → SHELF 作证据。",
        "transition": "This boundary gives us a clear recommendation about what LEGO should own and what it should outsource.",
        "remote": "CLICK once after the transition.",
        "open": "Entertainment is related diversification only when value returns to the brick.",
        "anchors": "shared resources + buildable return = related",
        "number": "One boundary: no return to play, no strategic fit",
        "close": "Screen-to-shelf proves the return loop.",
    },
    {
        "num": 16,
        "title": "Recommendation",
        "time": "0:40",
        "clicks": "1 click",
        "purpose": "把全部分析收束为明确、可执行的所有权与伙伴策略。",
        "script": (
            "My recommendation is not for LEGO to retreat from entertainment. LEGO should expand selectively as a hybrid ecosystem orchestrator. "
            "It should own more worlds: characters, story universes, digital spaces, communities, and play ideas that can become products. But it should "
            "own fewer studios. Lucasfilm, film studios, Epic Games, and other specialists can supply production and platform capabilities. LEGO must retain "
            "control of the brand, play principles, data, child safety, and the connection to physical play. Partner for production. Control the play system."
        ),
        "delivery": "左右两栏保持平衡：先 own more worlds，再 own fewer studios；最后两句逐句停顿。",
        "transition": "This leads to the final answer to our strategic question.",
        "remote": "CLICK once after the transition.",
        "open": "LEGO should expand entertainment selectively, not retreat from it.",
        "anchors": "own worlds / fewer studios / keep play control",
        "number": "One hybrid orchestrator model",
        "close": "Partner for production. Control the play system.",
    },
    {
        "num": 17,
        "title": "The Last Brick",
        "time": "0:30",
        "clicks": "1 click",
        "purpose": "用一个可记忆问题取代抽象结论。",
        "script": (
            "LEGO should not try to become a conventional entertainment conglomerate. Its stronger future is to keep the brick at the center while "
            "characters, licensed stories, owned worlds, and live platforms orbit around it. That keeps entertainment ambitious without losing the source "
            "of LEGO's advantage. The final decision rule is simple: will the next entertainment investment make a physical brick more valuable?"
        ),
        "delivery": "先让观众看中央积木两秒，再扫四个 orbit 标签。最后问题要放慢，不立即按键。",
        "transition": "We can turn that principle into a practical three-question investment test.",
        "remote": "CLICK once after the transition.",
        "open": "LEGO's future is not to become an entertainment conglomerate.",
        "anchors": "brick at center; media in orbit",
        "number": "One final decision rule",
        "close": "Will this investment make a physical brick more valuable?",
    },
    {
        "num": 18,
        "title": "The Brick Test",
        "time": "0:40",
        "clicks": "1 click",
        "purpose": "把推荐转化为能真正筛选投资项目的三道门。",
        "script": (
            "Every entertainment proposal should pass the Brick Test. First, core fit: does it use LEGO's system, brand, visual grammar, or world-building "
            "capability? Second, play return: can its characters, locations, or mechanics become new physical building experiences? Third, control: can LEGO "
            "own the play rules while specialist partners provide production? A project needs three strong yeses. If all three answers are yes, LEGO should build the world."
        ),
        "delivery": "逐项指 CORE FIT、PLAY RETURN、CONTROL。说 three strong yeses 时手势收拢到黄色结论。",
        "transition": "Now let us apply that logic creatively.",
        "remote": "CLICK once after the transition.",
        "open": "Every entertainment proposal should pass the Brick Test.",
        "anchors": "core fit / play return / control",
        "number": "3× YES",
        "close": "Three yeses—build the world.",
    },
    {
        "num": 19,
        "title": "Audience Challenge",
        "time": "0:45",
        "clicks": "1 click",
        "purpose": "让观众用 LEGO 的逻辑创造世界，而不是只被动听结论。",
        "script": (
            "Imagine that you control LEGO's next entertainment world. Give it one memorable character, one clear conflict, and one place that audiences "
            "would immediately want to build. Then ask: what new form of play does this combination create? This is the difference between simply producing "
            "content and creating a LEGO world. A good LEGO story does not end on the screen. It gives audiences a reason to build, rebuild, and continue the "
            "experience themselves. Character plus conflict plus a buildable place creates new play."
        ),
        "delivery": "问完后停 3–5 秒。若时间允许请一位观众用一句话回答；没人回答就直接说公式。",
        "transition": "The final theater shows how LEGO repeated this pattern across four generations.",
        "remote": "CLICK once only after the audience pause and transition.",
        "open": "If you controlled LEGO's next world, what would you build?",
        "anchors": "character + conflict + buildable place → new play",
        "number": "Pause 3–5 seconds",
        "close": "A good LEGO story gives audiences a reason to build.",
    },
    {
        "num": 20,
        "title": "Story Theater",
        "time": "0:40",
        "clicks": "Do not click",
        "purpose": "按舞台实际从左到右的顺序复盘：角色、授权、拥有世界、伙伴扩张。",
        "script": (
            "This theater brings the whole journey together. In 1978, LEGO gave the brick system characters. In 1999, LEGO borrowed a famous story and "
            "learned how to translate fandom into buildable play. In 2011, NINJAGO proved LEGO could own a renewable world. In 2023, LEGO Fortnite showed "
            "how a partner could help that world scale into persistent digital participation. The media changed, but the strategic loop remained consistent: "
            "create a world, invite participation, and return that energy to physical building. LEGO did not leave the brick behind. It learned how to make one "
            "brick the entrance to countless story worlds. Thank you."
        ),
        "delivery": "严格按左到右 1978 → 1999 → 2011 → 2023。‘Thank you’ 后看观众，不再按遥控器。",
        "transition": "End. Hold the final theater on screen.",
        "remote": "DO NOT CLICK. Another click returns to the crew.",
        "open": "This theater brings the whole journey together.",
        "anchors": "give characters → borrow story → own world → partner to scale",
        "number": "1978 / 1999 / 2011 / 2023",
        "close": "One brick became the entrance to countless story worlds. Thank you.",
    },
]


SOURCES = [
    "The LEGO Group History — Role Play / Minifigure timeline: https://www.lego.com/nl-nl/history/articles/f-role-play",
    "The LEGO Group — LEGO Star Wars celebrates 25 years (2024): https://www.lego.com/en-us/aboutus/news/2024/january/lego-star-wars-25-anniversary",
    "The LEGO Group Annual Report 2004: https://www.lego.com/cdn/cs/aboutus/assets/blt07abb4b8a3da3f39/Annual_Report_2004_ENG.pdf",
    "The LEGO Group Annual Report 2005: https://www.lego.com/cdn/cs/aboutus/assets/blt6eacf5a8b7af1359/Annual_Report_2005_ENG.pdf",
    "The LEGO Group — NINJAGO 15th Anniversary (2026): https://www.lego.com/en-gb/aboutus/news/2026/january/lego-ninjago-15th-anniversary",
    "The LEGO Group — 2014 Annual Results: https://www.lego.com/da-dk/aboutus/news/2019/november/lego-2014-annual-results",
    "The LEGO Group and Universal Pictures film partnership (2020): https://www.lego.com/en-us/aboutus/news/2020/april/universal-pictures-and-lego-group",
    "The LEGO Group — LEGO Fortnite launch (2023): https://www.lego.com/en-gb/aboutus/news/2023/december/the-adventure-is-building-lego-fortnite-is-live-",
    "The LEGO Group — LEGO Fortnite physical sets (2024): https://www.lego.com/en-gb/aboutus/news/2024/july/lego-fortnite-launch-2024?locale=en-gb",
    "The LEGO Group — 2024 Annual Results (87M+ players): https://www.lego.com/en-us/aboutus/news/2025/march/lego-group-delivers-record-results-in-2024",
]


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, *, bottom: str | None = None, left: str | None = None, size: int = 10) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge, color in (("bottom", bottom), ("left", left)):
        if color:
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), str(size))
            node.set(qn("w:space"), "4")
            node.set(qn("w:color"), color)
            p_bdr.append(node)


def set_run_font(run, *, size: float | None = None, color: str = INK, bold: bool | None = None, italic: bool | None = None, font: str = "Calibri") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.color.rgb = rgb(color)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_bullet_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    existing = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    existing_nums = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = max(existing_nums, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    p_pr.append(num_pr)


def configure_document(doc: Document, running_label: str) -> int:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("BUILDING BEYOND THE BRICK")
    set_run_font(run, size=8.5, color=LEGO_RED, bold=True)
    run = hp.add_run(f"    |    {running_label}")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_field(fp)

    doc.core_properties.title = running_label
    doc.core_properties.subject = "LEGO entertainment strategy presentation"
    doc.core_properties.author = "Harvey Yang"
    return add_bullet_numbering(doc)


def add_bullet(doc: Document, text: str, num_id: int, *, size: float = 10.5, color: str = INK) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    apply_numbering(p, num_id)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)


def add_cover(doc: Document, *, kicker: str, title: str, subtitle: str, metrics: str, core_story: str) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(kicker.upper())
    set_run_font(run, size=10, color=LEGO_RED, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    set_run_font(run, size=29, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(subtitle)
    set_run_font(run, size=14, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.right_indent = Inches(0.35)
    set_paragraph_shading(p, PALE_YELLOW)
    set_paragraph_border(p, left=LEGO_YELLOW, size=18)
    run = p.add_run(core_story)
    set_run_font(run, size=13, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    set_paragraph_shading(p, INK)
    run = p.add_run(metrics)
    set_run_font(run, size=10.5, color=LEGO_YELLOW, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared for live click-through delivery")
    set_run_font(run, size=10, color=MUTED, italic=True)
    doc.add_page_break()


def add_label_paragraph(doc: Document, label: str, text: str, *, fill: str | None = None, color: str = INK, label_color: str = DARK_BLUE, size: float = 10.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.2
    if fill:
        set_paragraph_shading(p, fill)
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
    run = p.add_run(f"{label}: ")
    set_run_font(run, size=size, color=label_color, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)


def add_slide_block(doc: Document, slide: dict, *, compact: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0 if compact else 3)
    p.paragraph_format.space_after = Pt(4 if compact else 5)
    p.paragraph_format.keep_with_next = True
    set_paragraph_shading(p, INK)
    run = p.add_run(f"SLIDE {slide['num']:02d}   |   {slide['time']}   |   {slide['clicks'].upper()}")
    set_run_font(run, size=9 if compact else 9.5, color=LEGO_YELLOW, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5 if compact else 7)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(slide["title"])
    set_run_font(run, size=13 if compact else 15, color=LEGO_RED, bold=True)

    if compact:
        add_label_paragraph(doc, "OPEN", slide["open"], size=9.3)
        add_label_paragraph(doc, "ANCHORS", slide["anchors"], size=9.3)
        add_label_paragraph(doc, "NUMBER", slide["number"], size=9.3)
        add_label_paragraph(doc, "CLOSE", slide["close"], size=9.3, fill=PALE_YELLOW, label_color=LEGO_RED)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        set_paragraph_border(p, bottom=PALE_BLUE, size=8)
        run = p.add_run(slide["remote"])
        set_run_font(run, size=8.7, color=LEGO_BLUE, bold=True)
        return

    add_label_paragraph(doc, "本页任务", slide["purpose"], size=9.5, color=MUTED, label_color=LEGO_RED)
    add_label_paragraph(doc, "SAY", slide["script"], size=10.5, label_color=LEGO_BLUE)
    add_label_paragraph(doc, "动作与节奏", slide["delivery"], size=9.4, color=MUTED, label_color=LEGO_GREEN)
    add_label_paragraph(doc, "TRANSITION", slide["transition"], size=10, fill=PALE_YELLOW, label_color=LEGO_RED)
    add_label_paragraph(doc, "REMOTE", slide["remote"], size=9.3, color=LEGO_BLUE, label_color=LEGO_BLUE)
    if slide.get("source"):
        add_label_paragraph(doc, "Source note", slide["source"], size=8.2, color=MUTED, label_color=MUTED)


def build_detailed() -> None:
    doc = Document()
    bullet_num_id = configure_document(doc, "Detailed Speaker Walkthrough")
    add_cover(
        doc,
        kicker="Presentation Walkthrough",
        title="Building Beyond the Brick",
        subtitle="A page-by-page English speaker script for the LEGO entertainment strategy story",
        metrics="20 SLIDES   |   14:00 TARGET   |   CLICK-THROUGH DELIVERY",
        core_story="LEGO did not abandon the brick to become an entertainment company. It used entertainment to make the brick system more valuable.",
    )

    doc.add_heading("How to use this walkthrough", level=1)
    add_bullet(doc, "Do not read the slide. Each page answers one question; the script supplies the causal link to the next page.", bullet_num_id)
    add_bullet(doc, "Text inside [BRACKETS] is an action, pause, or remote cue and should not be spoken aloud.", bullet_num_id)
    add_bullet(doc, "Aim for 130–145 words per minute. Protect the transition sentence even if you shorten the body.", bullet_num_id)
    add_bullet(doc, "On evidence-heavy pages, state the caveat exactly: contribution is not the same as single-cause proof.", bullet_num_id)
    add_bullet(doc, "The web deck advances with click-anywhere, Right Arrow, or Page Down. Avoid clicking video or media controls.", bullet_num_id)

    doc.add_heading("Delivery spine", level=2)
    add_label_paragraph(doc, "1", "The brick system already contained the grammar of stories.", fill=LIGHT, label_color=LEGO_RED)
    add_label_paragraph(doc, "2", "Licensing taught LEGO translation; owned worlds created continuity; platforms created recurrence.", fill=LIGHT, label_color=LEGO_RED)
    add_label_paragraph(doc, "3", "The 2004 crisis set the boundary: control the core, partner for specialist execution.", fill=LIGHT, label_color=LEGO_RED)
    add_label_paragraph(doc, "4", "Every entertainment investment must return value to physical play.", fill=PALE_YELLOW, label_color=LEGO_RED)
    doc.add_page_break()

    for i, slide in enumerate(SLIDES):
        add_slide_block(doc, slide)
        if slide["num"] == 3:
            doc.add_page_break()
        elif slide["num"] < 3 and slide["num"] % 2 == 0:
            doc.add_page_break()
        elif slide["num"] >= 4 and slide["num"] < 20 and slide["num"] % 2 == 1:
            doc.add_page_break()
        elif slide["num"] == 20:
            doc.add_page_break()

    doc.add_heading("Research sources used in the script", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Use these for fact-checking and Q&A. They are not intended to be read aloud.")
    set_run_font(run, size=10, color=MUTED, italic=True)
    for source in SOURCES:
        add_bullet(doc, source, bullet_num_id, size=9.2, color=MUTED)

    doc.save(DETAILED_PATH)


def build_cue_sheet() -> None:
    doc = Document()
    bullet_num_id = configure_document(doc, "Presenter Cue Sheet")
    add_cover(
        doc,
        kicker="Presenter Cue Sheet",
        title="Building Beyond the Brick",
        subtitle="Short prompts, numbers, closing lines, and remote cues",
        metrics="20 SLIDES   |   4 SLIDES PER PAGE   |   14:00 TARGET",
        core_story="If you lose your place, return to this sentence: entertainment must make the brick system more valuable.",
    )

    del bullet_num_id  # The cue sheet intentionally uses labeled prompt blocks instead of lists.
    for i, slide in enumerate(SLIDES):
        add_slide_block(doc, slide, compact=True)
        if slide["num"] % 4 == 0 and slide["num"] < 20:
            doc.add_page_break()

    doc.save(CUE_PATH)


if __name__ == "__main__":
    build_detailed()
    build_cue_sheet()
    print(DETAILED_PATH)
    print(CUE_PATH)
